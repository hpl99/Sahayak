"""
Beneficiary Resume Generator for Voice for Livelihood.

Generates a structured, one-page DOCX resume from the saved beneficiary profile
and recommended NSQF skilling pathways.
"""

import io
from datetime import datetime
from typing import Any, Dict, List, Optional
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_resume_docx(
    profile: Dict[str, Any],
    recommendations: Optional[List[Dict[str, Any]]] = None,
) -> bytes:
    """
    Generate a 1-page structured .docx resume for a beneficiary.

    Returns raw byte data suitable for st.download_button or file saving.
    """
    doc = docx.Document()

    # Set 0.5 inch margins to ensure it fits neatly on one page
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # 1. Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("SKILL INDIA — BENEFICIARY LIVELIHOOD RESUME")
    title_run.bold = True
    title_run.font.size = Pt(15)
    title_run.font.color.rgb = RGBColor(26, 54, 93)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run(
        f"Beneficiary ID: {profile.get('beneficiary_id', 'N/A')}  |  "
        f"District: {profile.get('district', 'N/A')}  |  "
        f"Date: {datetime.now().strftime('%d-%b-%Y')}"
    )
    sub_run.font.size = Pt(9.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(74, 85, 104)

    # 2. Personal & Background Summary
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Beneficiary Profile & Personal Background")
    h1_run.bold = True
    h1_run.font.size = Pt(11.5)
    h1_run.font.color.rgb = RGBColor(43, 108, 176)

    t1 = doc.add_table(rows=4, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = True

    t1.cell(0, 0).text = f"Full Name: {profile.get('name', 'N/A')}"
    t1.cell(0, 1).text = f"Age / Gender: {profile.get('age', 'N/A')} / {profile.get('gender', 'N/A')}"

    t1.cell(1, 0).text = f"Education Level: {profile.get('education_level', 'N/A')}"
    t1.cell(1, 1).text = f"Preferred Language: {profile.get('language', 'Hindi')}"

    t1.cell(2, 0).text = f"Current Livelihood: {profile.get('current_livelihood') or 'Informal / Unemployed'}"
    t1.cell(2, 1).text = f"Family Occupation: {profile.get('family_occupation') or 'Not specified'}"

    t1.cell(3, 0).text = f"Employment Preference: {profile.get('employment_preference', 'N/A')}"
    t1.cell(3, 1).text = f"Mobility: {profile.get('mobility_constraints', 'Local only')}"

    # Style table text
    for row in t1.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # 3. Stated Skills & Prior Experience
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Stated Skills, Experience & Aspirations")
    h2_run.bold = True
    h2_run.font.size = Pt(11.5)
    h2_run.font.color.rgb = RGBColor(43, 108, 176)

    skills_text = profile.get("skills", "").strip() or "General work readiness / informal tasks"
    p_skills = doc.add_paragraph()
    r_sk_label = p_skills.add_run("• Prior Skills & Experience: ")
    r_sk_label.bold = True
    r_sk_label.font.size = Pt(9.5)
    r_sk_val = p_skills.add_run(skills_text)
    r_sk_val.font.size = Pt(9.5)

    interests_text = profile.get("interests", "").strip() or "Skilling and sustainable livelihood"
    p_int = doc.add_paragraph()
    r_in_label = p_int.add_run("• Learning Goals & Interests: ")
    r_in_label.bold = True
    r_in_label.font.size = Pt(9.5)
    r_in_val = p_int.add_run(interests_text)
    r_in_val.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # 4. Recommended NSQF Skilling Pathways
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Recommended NSQF Skilling & Livelihood Pathways")
    h3_run.bold = True
    h3_run.font.size = Pt(11.5)
    h3_run.font.color.rgb = RGBColor(43, 108, 176)

    if recommendations:
        for idx, rec in enumerate(recommendations[:3], 1):
            p_rec = doc.add_paragraph()
            r_t = p_rec.add_run(
                f"{idx}. {rec.get('trade_name', '')} "
                f"(NSQF Level {rec.get('nsqf_level', '4')} · {rec.get('sector', '')})"
            )
            r_t.bold = True
            r_t.font.size = Pt(9.5)

            p_det = doc.add_paragraph()
            r_det = p_det.add_run(
                f"   Local Demand: {rec.get('demand_score', 0):.0f}/10  |  "
                f"Avg Monthly Wage: ₹{rec.get('avg_monthly_wage_inr', 0):,}  |  "
                f"Match Score: {rec.get('score', 0):.1f}"
            )
            r_det.font.size = Pt(9)
            r_det.font.color.rgb = RGBColor(100, 116, 139)

            if rec.get("explanations"):
                for exp in rec["explanations"]:
                    p_exp = doc.add_paragraph()
                    r_exp = p_exp.add_run(f"   {exp}")
                    r_exp.font.size = Pt(8.5)
                    r_exp.font.italic = True
    else:
        assigned_trade = profile.get("recommended_trade") or "General NSQF Skilling Assessment"
        p_rec = doc.add_paragraph()
        r_t = p_rec.add_run(f"• Assigned Pathway: {assigned_trade}")
        r_t.bold = True
        r_t.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # 5. Training Status & Verification Disclaimer
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Training Status & Verification")
    h4_run.bold = True
    h4_run.font.size = Pt(11.5)
    h4_run.font.color.rgb = RGBColor(43, 108, 176)

    p_status = doc.add_paragraph()
    r_st = p_status.add_run(f"Current Training Status: {profile.get('training_status', 'Not Started')}")
    r_st.bold = True
    r_st.font.size = Pt(9.5)

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_note = p_note.add_run(
        "Generated via Voice for Livelihood Prototype (SIH PS 26097) — "
        "Verified with local demand analytics."
    )
    r_note.font.size = Pt(8)
    r_note.font.italic = True
    r_note.font.color.rgb = RGBColor(148, 163, 184)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_resume_preview_text(
    profile: Dict[str, Any],
    recommendations: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """Generate clean plain text summary for quick inspection."""
    lines = [
        "=" * 60,
        f"SKILL INDIA — BENEFICIARY RESUME: {profile.get('name', 'N/A')}",
        f"ID: {profile.get('beneficiary_id', 'N/A')} | District: {profile.get('district', 'N/A')}",
        "=" * 60,
        f"Age / Gender: {profile.get('age')} / {profile.get('gender')}",
        f"Education: {profile.get('education_level')}",
        f"Language: {profile.get('language')}",
        f"Current Work: {profile.get('current_livelihood') or 'None'}",
        f"Mobility: {profile.get('mobility_constraints')}",
        f"Employment Preference: {profile.get('employment_preference')}",
        "-" * 60,
        f"Skills: {profile.get('skills') or 'None'}",
        f"Interests: {profile.get('interests') or 'None'}",
        "-" * 60,
        "Recommended NSQF Pathways:",
    ]
    if recommendations:
        for idx, rec in enumerate(recommendations[:3], 1):
            lines.append(
                f"{idx}. {rec.get('trade_name')} (NSQF L{rec.get('nsqf_level')}) - "
                f"₹{rec.get('avg_monthly_wage_inr'):,}/mo (Demand: {rec.get('demand_score'):.0f}/10)"
            )
            for exp in rec.get("explanations", []):
                lines.append(f"   {exp}")
    else:
        lines.append(f"• Assigned Trade: {profile.get('recommended_trade') or 'Pending'}")

    lines.append("=" * 60)
    return "\n".join(lines)
